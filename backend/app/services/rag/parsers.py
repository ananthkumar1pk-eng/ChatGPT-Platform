"""
Multi-Format Document Parsing Engine.
Supports: PDF, DOCX, TXT, CSV, JSON files with per-page / per-section attribution.
"""

import json
import os
from typing import List, Dict, Any, Tuple
import pypdf
import docx
import pandas as pd


class DocumentParser:
    @staticmethod
    def parse_file(file_path: str, filename: str) -> Tuple[List[Dict[str, Any]], int]:
        """
        Parses a document into a structured list of page/section items.
        Returns (list_of_pages, total_pages).
        Each page item format: {"page_number": int, "text": str, "meta": dict}
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentParser._parse_docx(file_path)
        elif ext == ".csv":
            return DocumentParser._parse_csv(file_path)
        elif ext == ".json":
            return DocumentParser._parse_json(file_path)
        elif ext in [".txt", ".md", ".markdown", ".log", ".py", ".js", ".ts", ".html"]:
            return DocumentParser._parse_txt(file_path)
        else:
            # Default text fallback
            return DocumentParser._parse_txt(file_path)

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[List[Dict[str, Any]], int]:
        pages_out = []
        try:
            reader = pypdf.PdfReader(file_path)
            total = len(reader.pages)
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_out.append({
                        "page_number": idx + 1,
                        "text": text.strip(),
                        "meta": {"type": "pdf_page"}
                    })
            if not pages_out and total > 0:
                pages_out.append({
                    "page_number": 1,
                    "text": (
                        f"[Document Notice: This PDF file contains {total} pages of scanned raster images or slides "
                        f"without an embedded selectable text layer. To perform rich semantic search and question-answering, "
                        f"please upload a text-based PDF, DOCX, TXT, CSV, or JSON document.]"
                    ),
                    "meta": {"type": "pdf_scanned_notice"}
                })
            return pages_out, max(1, total)
        except Exception as e:
            # Fallback if corrupted
            return [{"page_number": 1, "text": f"[Error reading PDF: {str(e)}]", "meta": {}}], 1

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[List[Dict[str, Any]], int]:
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if row_text:
                        paragraphs.append(f"[Table Row] {row_text}")

            full_text = "\n\n".join(paragraphs)
            # Estimate pages (~400 words per page)
            words = full_text.split()
            page_count = max(1, len(words) // 400 + (1 if len(words) % 400 else 0))

            pages_out = []
            words_per_page = max(1, len(words) // page_count + 1)
            for i in range(page_count):
                chunk_words = words[i * words_per_page : (i + 1) * words_per_page]
                if chunk_words:
                    pages_out.append({
                        "page_number": i + 1,
                        "text": " ".join(chunk_words),
                        "meta": {"type": "docx_section"}
                    })

            return pages_out, page_count
        except Exception as e:
            return [{"page_number": 1, "text": f"[Error reading DOCX: {str(e)}]", "meta": {}}], 1

    @staticmethod
    def _parse_csv(file_path: str) -> Tuple[List[Dict[str, Any]], int]:
        try:
            df = pd.read_csv(file_path)
            # Create a structured text overview + row blocks
            overview = f"CSV Dataset Summary: {len(df)} rows, {len(df.columns)} columns. Columns: {', '.join(df.columns.astype(str))}\n\n"

            row_chunks = []
            batch_size = 25  # 25 rows per section
            num_batches = max(1, len(df) // batch_size + (1 if len(df) % batch_size else 0))

            for b in range(num_batches):
                subset = df.iloc[b * batch_size : (b + 1) * batch_size]
                text_block = overview if b == 0 else ""
                text_block += subset.to_string(index=True)
                row_chunks.append({
                    "page_number": b + 1,
                    "text": text_block,
                    "meta": {"type": "csv_batch", "rows": f"{b*batch_size} - {min(len(df), (b+1)*batch_size)}"}
                })

            return row_chunks, num_batches
        except Exception as e:
            return [{"page_number": 1, "text": f"[Error reading CSV: {str(e)}]", "meta": {}}], 1

    @staticmethod
    def _parse_json(file_path: str) -> Tuple[List[Dict[str, Any]], int]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)

            pretty_text = json.dumps(data, indent=2)
            lines = pretty_text.splitlines()
            lines_per_page = 100
            total_pages = max(1, len(lines) // lines_per_page + (1 if len(lines) % lines_per_page else 0))

            pages_out = []
            for i in range(total_pages):
                chunk_lines = lines[i * lines_per_page : (i + 1) * lines_per_page]
                pages_out.append({
                    "page_number": i + 1,
                    "text": "\n".join(chunk_lines),
                    "meta": {"type": "json_block"}
                })

            return pages_out, total_pages
        except Exception as e:
            return [{"page_number": 1, "text": f"[Error reading JSON: {str(e)}]", "meta": {}}], 1

    @staticmethod
    def _parse_txt(file_path: str) -> Tuple[List[Dict[str, Any]], int]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

            words = text.split()
            words_per_page = 400
            total_pages = max(1, len(words) // words_per_page + (1 if len(words) % words_per_page else 0))

            pages_out = []
            for i in range(total_pages):
                chunk_words = words[i * words_per_page : (i + 1) * words_per_page]
                if chunk_words:
                    pages_out.append({
                        "page_number": i + 1,
                        "text": " ".join(chunk_words),
                        "meta": {"type": "txt_page"}
                    })

            return pages_out if pages_out else [{"page_number": 1, "text": text, "meta": {}}], total_pages
        except Exception as e:
            return [{"page_number": 1, "text": f"[Error reading text: {str(e)}]", "meta": {}}], 1
